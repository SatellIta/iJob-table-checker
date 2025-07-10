from docx import Document
from .conference_info import Conference
from UI.some_dialog import FailFileMessage
import os, re

def open_document(file_path):
    '''
    读取word文件,并返回文档对象

    args:
        file_path: word文档的路径
    '''
    try:
        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
            
        # 检查文件扩展名
        if not file_path.lower().endswith('.docx'):
            raise ValueError("不是有效的Word文档(.docx)文件")
            
        # 尝试打开文档
        doc = Document(file_path)
        return doc
        
    except Exception as e:
        
        oops = FailFileMessage(None, str(os.path.basename(file_path)), str(e))
        oops.exec()
        return None

# 传入doc文件,创建并返回所有包含所有宣讲会的列表
def storage_info(doc):
    doc = remove_headers(doc)   # 清除页眉
    row_data = read_document(doc)  # 读取行数据
    grouped_data = group_by_conference(row_data)   # 自动查找宣讲会个数n，将行数据改写成n个字典
    conferences = []
    
    for index, text in grouped_data.items():
        conf = create_conference(index, text)
        if conf:    
            conferences.append(conf)
    
    return conferences

# 清除页眉
def remove_headers(doc):
    """
    删除Word文档中的所有页眉
    
    args:
        doc: Document对象
    """
    for section in doc.sections:
        # 断开页眉与上一节的链接
        section.header.is_linked_to_previous = False
        
        # 清空页眉中的所有段落
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                run.text = ""
    
    return doc

# 读取doc中所有行信息的函数
def read_document(doc):
    row_data = []
    if doc.paragraphs:
        for line in doc.paragraphs:
            #line.text.strip()用于消除一行中开头和结尾多余的空白字符,这里起判断是否为空行的作用
            if line.text.strip():       
                row_data.append(line)
            else:
                continue
    else:
        print(doc)

    return row_data

# 自动查找宣讲会个数n，将行数据改写成n个字典,每个字典包含该宣讲会所有的行
def group_by_conference(row_data):
    grouped_data = {}
    pattern = re.compile(r'No\.\s*(\d+)')  # 匹配"No.数字"模式
    current_index = None

    for row in row_data:
        text = row.text.strip()  #这里是清除多余空白字符
        match = pattern.search(text)

        if match:
            current_index = int(match.group(1)) # 提取捕获组，得到正确的宣讲会编号
            if current_index not in grouped_data:
                grouped_data[current_index] = []

        if current_index is not None:
            grouped_data[current_index].append(text)

    return grouped_data


def create_conference(index, text):
    """
    从收集的文本行创建Conference对象
    
    参数:
        index: 会议编号
        text: 会议相关的文本行列表
    """
    if not text:
        return None
    
    # 解析标题
    title = text[0]  
    
    # 尝试解析地点、日期和时间
    place = ""
    date = ""
    time = ""
    detail = ""
    
    # 寻找含有地点的行
    place_pattern = re.compile(r'宣讲地点[:：]\s*(.*)', re.IGNORECASE)
    # 寻找含有日期的行
    date_pattern = re.compile(r'\d{4}-\d{2}-\d{2}', re.IGNORECASE)
    # 寻找含有时间的行
    time_pattern = re.compile(r'\d{2}:\d{2}-\d{2}:\d{2}', re.IGNORECASE)
    # 寻找含有详情的行
    detail_pattern = re.compile(r'宣讲详情[：:]\s*(.*)')
    
    for line in text[1:]:  # 跳过标题行
        # 检查地点
        place_match = place_pattern.search(line)
        if place_match:
            place = place_match.group(1).strip()
            
        # 检查日期
        date_match = date_pattern.search(line)
        if date_match:
            date = date_match.group(0).strip()
            
        # 检查时间
        time_match = time_pattern.search(line)
        if time_match:
            time = time_match.group(0).strip()
            
        # 其他行作为详情
        detail_match = detail_pattern.search(str(text))
        if detail_match:
            detail = detail_match.group(1).strip()
            detail = detail[ :-2]    # 去掉列表强转list的最后的"']"字符串

    # 创建并返回Conference对象
    return Conference(
        index=index,
        title=title.replace(f"No.{index}", "").strip(),  # 移除标题中的编号
        place=place,
        date=date,
        time=time,
        detail=detail
    )
        